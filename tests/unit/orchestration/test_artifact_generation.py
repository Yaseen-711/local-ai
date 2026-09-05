"""Unit tests for Deterministic Artifact Generation capability and generators."""

from pathlib import Path
import pytest
import openpyxl
import docx

from orchestration.capabilities.base import CapabilityContext
from orchestration.capabilities.builtin.artifact import (
    ArtifactFormat,
    ArtifactGenerationCapability,
    ArtifactGenerationRequest,
    DocxGenerator,
    PdfGenerator,
    XlsxGenerator,
)
from orchestration.domain.references import ArtifactReference


@pytest.fixture
def temp_output_dir(tmp_path: Path) -> Path:
    out_dir = tmp_path / "artifacts"
    out_dir.mkdir(parents=True)
    return out_dir


def test_xlsx_generator_list_of_dicts(temp_output_dir: Path):
    """Verify XlsxGenerator converts a list of dictionaries into an Excel sheet."""
    target = temp_output_dir / "employees.xlsx"
    data = [
        {"ID": 101, "Name": "Alice Smith", "Department": "Engineering", "Salary": 120000},
        {"ID": 102, "Name": "Bob Jones", "Department": "Design", "Salary": 95000},
    ]

    req = ArtifactGenerationRequest(
        format=ArtifactFormat.XLSX,
        filename="employees.xlsx",
        title="Employee Directory",
        data=data,
    )
    XlsxGenerator.generate(req, target)

    assert target.exists()
    wb = openpyxl.load_workbook(target)
    ws = wb.active
    assert ws.title == "Employee Directory"[:31]
    # Header row
    headers = [cell.value for cell in ws[1]]
    assert headers == ["ID", "Name", "Department", "Salary"]
    # Row 1
    row1 = [cell.value for cell in ws[2]]
    assert row1 == [101, "Alice Smith", "Engineering", 120000]


def test_xlsx_generator_multi_sheets(temp_output_dir: Path):
    """Verify XlsxGenerator handles multi-sheet dictionary input."""
    target = temp_output_dir / "multi_sheet.xlsx"
    data = {
        "Q1_Sales": [["Region", "Amount"], ["North", 50000], ["South", 42000]],
        "Q2_Sales": [["Region", "Amount"], ["North", 58000], ["South", 49000]],
    }

    req = ArtifactGenerationRequest(
        format=ArtifactFormat.XLSX,
        filename="multi_sheet.xlsx",
        data=data,
    )
    XlsxGenerator.generate(req, target)

    assert target.exists()
    wb = openpyxl.load_workbook(target)
    sheet_names = wb.sheetnames
    assert "Q1_Sales" in sheet_names
    assert "Q2_Sales" in sheet_names


def test_docx_generator_prose_and_table(temp_output_dir: Path):
    """Verify DocxGenerator compiles title, markdown prose, bullet points, and tables."""
    target = temp_output_dir / "report.docx"
    content = """
## Executive Summary
This document provides quarterly revenue and project updates.

### Key Milestones
- Completed milestone 1 ahead of schedule
- Began architecture design for milestone 2
"""
    table_data = [
        ["Phase", "Status", "Owner"],
        ["Inception", "Done", "Alice"],
        ["Implementation", "In Progress", "Bob"],
    ]

    req = ArtifactGenerationRequest(
        format=ArtifactFormat.DOCX,
        filename="report.docx",
        title="Quarterly Review",
        content=content,
        data=table_data,
    )
    DocxGenerator.generate(req, target)

    assert target.exists()
    doc = docx.Document(target)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    assert "Quarterly Review" in paragraphs
    assert "Executive Summary" in paragraphs
    assert any("Completed milestone 1" in p for p in paragraphs)
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "Phase"
    assert doc.tables[0].cell(1, 1).text == "Done"


def test_pdf_generator_prose_and_table(temp_output_dir: Path):
    """Verify PdfGenerator generates valid PDF with ReportLab."""
    target = temp_output_dir / "summary.pdf"
    content = "This is a detailed analysis paragraph explaining financial results."
    table_data = [
        ["Metric", "Value"],
        ["ARR", "$12.5M"],
        ["Net Margin", "22%"],
    ]

    req = ArtifactGenerationRequest(
        format=ArtifactFormat.PDF,
        filename="summary.pdf",
        title="Financial Status",
        content=content,
        data=table_data,
    )
    PdfGenerator.generate(req, target)

    assert target.exists()
    header_bytes = target.read_bytes()[:10]
    assert header_bytes.startswith(b"%PDF-")


def test_artifact_capability_xlsx_execution(temp_output_dir: Path):
    """Verify ArtifactGenerationCapability execution producing XLSX."""
    cap = ArtifactGenerationCapability(output_dir=temp_output_dir)
    assert cap.capability_id == "artifact.generate"

    ctx = CapabilityContext(execution_id="exec-art-1")
    result = cap.execute(
        parameters={"artifact_type": "xlsx", "filename": "budget.xlsx", "title": "2025 Budget"},
        inputs={"data": [["Item", "Cost"], ["Servers", 5000], ["Licenses", 2000]]},
        context=ctx,
    )

    assert result.output is not None
    assert result.output["name"] == "budget.xlsx"
    assert result.output["mime_type"] == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    assert len(result.output["sha256"]) == 64
    assert result.output["size_bytes"] > 0
    assert len(result.artifacts) == 1

    art = result.artifacts[0]
    assert isinstance(art, ArtifactReference)
    assert art.name == "budget.xlsx"
    assert art.size_bytes == result.output["size_bytes"]
    assert art.metadata["sha256"] == result.output["sha256"]


def test_artifact_capability_consume_document_tables(temp_output_dir: Path):
    """Verify ArtifactGenerationCapability seamlessly consumes tables from document.understand."""
    cap = ArtifactGenerationCapability(output_dir=temp_output_dir)
    ctx = CapabilityContext(execution_id="exec-art-2")

    # Structure emitted by document.understand (list of DocumentTable dicts)
    doc_tables = [
        {
            "table_id": "Revenue_Table",
            "page_number": 1,
            "num_rows": 3,
            "num_cols": 2,
            "grid": [["Year", "Revenue"], ["2023", "$10M"], ["2024", "$15M"]],
        }
    ]

    result = cap.execute(
        parameters={"artifact_type": "xlsx", "filename": "extracted_tables.xlsx"},
        inputs={"data": doc_tables},
        context=ctx,
    )

    assert result.output is not None
    assert (temp_output_dir / "extracted_tables.xlsx").exists()
    wb = openpyxl.load_workbook(temp_output_dir / "extracted_tables.xlsx")
    assert "Revenue_Table" in wb.sheetnames


def test_artifact_capability_docx_and_pdf(temp_output_dir: Path):
    """Verify capability generates DOCX and PDF formats."""
    cap = ArtifactGenerationCapability(output_dir=temp_output_dir)
    ctx = CapabilityContext(execution_id="exec-art-3")

    # DOCX
    res_docx = cap.execute(
        parameters={"artifact_type": "docx", "filename": "doc1.docx", "title": "Doc Title"},
        inputs={"content": "Content paragraph"},
        context=ctx,
    )
    assert res_docx.artifacts[0].mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    assert (temp_output_dir / "doc1.docx").exists()

    # PDF
    res_pdf = cap.execute(
        parameters={"artifact_type": "pdf", "filename": "doc2.pdf", "title": "PDF Title"},
        inputs={"content": "PDF paragraph"},
        context=ctx,
    )
    assert res_pdf.artifacts[0].mime_type == "application/pdf"
    assert (temp_output_dir / "doc2.pdf").exists()


def test_artifact_capability_unsupported_format(temp_output_dir: Path):
    """Verify error on unsupported artifact type."""
    cap = ArtifactGenerationCapability(output_dir=temp_output_dir)
    ctx = CapabilityContext(execution_id="exec-art-4")

    with pytest.raises(ValueError, match="Unsupported artifact_type 'invalid_format'"):
        cap.execute(
            parameters={"artifact_type": "invalid_format"},
            inputs={},
            context=ctx,
        )
